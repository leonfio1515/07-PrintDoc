from openpyxl import load_workbook
from docx import Document
from docx.shared import Cm

import win32api
#--------------------------------------------------------------------#

excel = load_workbook('Uniformes-print.xlsx')
sheet = excel.active


word = Document()


section = word.sections[0]
section.page_width = Cm(8)   
section.left_margin = Cm(0.5)   
section.right_margin = Cm(0.5)  


for row in sheet.iter_rows(values_only=True):
    sucursal = row[0]
    num_fun = row[1]
    ci_fun = row[2]
    nombre = row[3]
    apellido = row[4]
    sexo = row[5]
    area = row[6]

    camisa = row[7]
    pantalon = row[8]
    abrigo = row[9]


    etiqueta = word.add_paragraph()
    etiqueta.add_run(f"Sucursal: {sucursal}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"Fun: {num_fun}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"CI: {ci_fun}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"Nombre: {nombre}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"Apllido: {apellido}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"Sexo: {sexo}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"Area: {area}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"Camisa: {camisa}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"Pantalon: {pantalon}")
    etiqueta.add_run("\n")
    etiqueta.add_run(f"Abrigo: {abrigo}")
    etiqueta.add_run("\n")
    word.add_page_break()

word.save('etiquetas.docx')


win32api.ShellExecute(0, "print", "etiquetas.docx", None, ".", 0)